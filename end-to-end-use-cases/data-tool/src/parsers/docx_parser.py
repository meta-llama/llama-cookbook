import os
import docx

class DOCXParser:
    def __init__(self):
        pass
    
    def parse(self, file_path):
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        doc = docx.Document(file_path)
        full_text = []
        elements = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():  # Skip empty paragraphs
                elements.append(('paragraph', i, para))
        table_indices = []
        last_para_idx = 0
        for i, para in enumerate(doc.paragraphs):
            last_para_idx = i
            
        for i, table in enumerate(doc.tables):
            approx_pos = last_para_idx + i + 1
            elements.append(('table', approx_pos, table))
        elements.sort(key=lambda x: x[1])
        
        for elem_type, _, elem in elements:
            if elem_type == 'paragraph':
                full_text.append(elem.text)
            elif elem_type == 'table':
                # Process the table
                table_text = self._process_table(elem)
                full_text.extend(table_text)
        
        return '\n'.join(full_text)
    #the logic for tables is a bit overkill but it works quite well
    def _process_table(self, table):
        table_lines = []
        
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            
            if row_texts:
                table_lines.append(' | '.join(row_texts))
        if table_lines:
            table_lines.append('')
            
        return table_lines
    
    def save(self, content, output_path):
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)